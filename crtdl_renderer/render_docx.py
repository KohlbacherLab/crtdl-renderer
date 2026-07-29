# Copyright 2026 Kohlbacher Lab, University of Tübingen
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
"""DOCX rendering via python-docx (optional dependency)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from .model import Query
from .render import (
    MUSTHAVE_NOTE,
    NO_CONSTRAINT,
    UNRESOLVED_NOTE,
    Row,
    block_formula,
    block_intro,
    block_rows,
    code_systems,
    cohort_rule,
    completeness_note,
    criteria_cards,
    criteria_head,
    group_attributes,
    group_filters,
    group_names,
    group_title,
    has_must_have,
    legend_for,
    unresolved_codes,
    version_line,
)


def _plain(text: str) -> str:
    """Strip the Markdown emphasis markers used in the shared sentence helpers."""
    return text.replace("**", "")


def _shade(paragraph, hex_colour: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    pPr.append(shd)


def _left_rail(paragraph, hex_colour: str, size: int) -> None:
    """A coloured left paragraph border — the DOCX equivalent of the accent rail."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_colour)
    borders.append(left)
    pPr.append(borders)


def _card_paragraphs(doc, cards, excl: bool, depth: int = 0) -> None:
    from docx.shared import Cm, Pt
    ramp = (["8A3B2C", "A8574A", "C2796D"] if excl else ["2C5F8A", "4A7BA8", "6D97C2"])
    rail = ramp[min(depth, 2)]
    for card in cards:
        if card.op:
            p = doc.add_paragraph()
            p.alignment = 1  # centred operator band between cards
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(card.op).bold = True
        head = doc.add_paragraph()
        head.paragraph_format.left_indent = Cm(0.45 * depth)
        head.paragraph_format.space_after = Pt(1)
        _left_rail(head, rail, 18 - 4 * min(depth, 2))
        if depth == 0:
            _shade(head, "FAF2F0" if excl else "F2F6FA")
        head.add_run(f"{card.label}  ").bold = True
        head.add_run(card.title)
        lines = card.constraints or ([NO_CONSTRAINT] if card.kind == "criterion" else [])
        for line in lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45 * depth + 0.4)
            p.paragraph_format.space_after = Pt(0)
            _left_rail(p, rail, 18 - 4 * min(depth, 2))
            p.add_run(f"· {line}").font.size = Pt(9)
        if card.children:
            _card_paragraphs(doc, card.children, excl, depth + 1)

def _criteria_table(doc, block, rows: list[Row]):
    from docx.shared import Cm, Pt
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for cell, h in zip(table.rows[0].cells, criteria_head(block), strict=False):
        cell.paragraphs[0].add_run(h).bold = True
    for r in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(r.number).bold = r.is_header
        p = cells[1].paragraphs[0]
        p.paragraph_format.left_indent = Cm(0.4 * r.indent)
        p.add_run(r.text).bold = r.is_header
        cp = cells[2].paragraphs[0]
        lines = r.constraints or ([] if r.is_header else [NO_CONSTRAINT])
        for i, line in enumerate(lines):
            (cp if i == 0 else cells[2].add_paragraph()).add_run(line).font.size = Pt(9)
    return table


def render_docx(q: Query, path: str | Path, today: date | None = None,
                layout: str = "cards") -> None:
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    doc.add_heading(f"Machbarkeitsanfrage: {q.display or q.source_name}", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Quelle: {q.source_name}   Erstellt: {(today or date.today()).isoformat()}"
                 + (f"   {version_line(q)}" if version_line(q) else "")).italic = True

    doc.add_heading("Kohortendefinition", level=1)
    doc.add_paragraph(_plain(cohort_rule(q)))
    for block, title in ((q.inclusion, "Einschlusskriterien"),
                         (q.exclusion, "Ausschlusskriterien")):
        if not block:
            continue
        doc.add_heading(title, level=2)
        doc.add_paragraph().add_run(_plain(block_intro(block))).bold = True
        note = doc.add_paragraph().add_run(
            f"Formale Struktur: {block_formula(block)} — {completeness_note(block)}")
        note.italic = True
        note.font.size = Pt(8)
        if layout == "table":
            _criteria_table(doc, block, block_rows(block))
        else:
            _card_paragraphs(doc, criteria_cards(block), block.kind == "exclusion")

    if q.attribute_groups:
        doc.add_heading("Datenextraktion", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for cell, h in zip(table.rows[0].cells, ["Modul / Profil", "Filter", "Attribute"], strict=False):
            cell.text = h
        names = group_names(q)
        for g in q.attribute_groups:
            cells = table.add_row().cells
            cells[0].text = group_title(g)
            cells[1].text = "\n".join(group_filters(g)) or "—"
            cells[2].text = "\n".join(group_attributes(g, names)) or "—"
        if has_must_have(q):
            doc.add_paragraph(MUSTHAVE_NOTE)

    systems = code_systems(q)
    doc.add_heading("Kodiersysteme", level=1)
    st = doc.add_table(rows=1, cols=3)
    st.style = "Light Grid Accent 1"
    for cell, h in zip(st.rows[0].cells, ["Kurzform", "System-URI", "Version(en)"], strict=False):
        cell.paragraphs[0].add_run(h).bold = True
    for label, uri, versions, ambiguous in systems:
        cells = st.add_row().cells
        cells[0].text = label + (" \u26a0" if ambiguous else "")
        cells[1].text = uri
        cells[2].text = versions
    if any(a for *_, a in systems):
        doc.add_paragraph("\u26a0 Diese Kurzform steht im Dokument für mehr als eine "
                          "System-URI. Die Kriterien nennen nur die Kurzform; maßgeblich "
                          "ist die URI aus dem Export.")

    doc.add_heading("Lesehilfe", level=1)
    legend = doc.add_table(rows=1, cols=2)
    legend.style = "Light Grid Accent 1"
    for cell, h in zip(legend.rows[0].cells, ["Notation", "Bedeutung"], strict=False):
        cell.paragraphs[0].add_run(h).bold = True
    for sym, meaning in legend_for(q):
        cells = legend.add_row().cells
        cells[0].text = sym
        cells[1].text = meaning

    if q.unresolved:
        doc.add_paragraph(UNRESOLVED_NOTE
                          + ", ".join(unresolved_codes(q))).runs[0].italic = True
    doc.save(str(path))
